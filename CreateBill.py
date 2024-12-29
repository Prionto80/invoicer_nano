from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from io import BytesIO


def create_bill(billNo, brand, showroom, address, mobile, board):
	#----Declaring Variables--------
	bill_no = billNo
	brand = brand
	showroom = showroom
	address = address
	mobile = mobile
	board = board
	#----Creating The document-------
	doc = Document()

	#document Heading--------------------------------------------------------
	heading = doc.add_paragraph(f"Bill No {bill_no}", style="Heading 1")
	doc.styles["Heading 1"].font.size = Pt(20)
	doc.styles["Heading 1"].font.name = "Calibri"
	heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	#------------------------------------------------------------------------



	#document Top Section-------------------------------------------------
	doc.styles["Normal"].font.size = Pt(11)
	doc.styles["Normal"].font.name = "Calibri"
	doc.styles["Heading 3"].font.size = Pt(11)
	doc.styles["Heading 3"].font.name = "Calibri"

	if brand == "m":
		brand1 = "Walton Digi-Tech Mobile"
	else:
		brand1 = "Walton Hi-Tech Industries (PLC) COM"

	paragraph_walton = doc.add_paragraph(f'\n{brand1}' , style="Heading 3")
	#paragraph_walton.bold = True
	paragraph_walton.paragraph_format.space_after = 0
	paragraph_shop = doc.add_paragraph(showroom)
	paragraph_shop.bold = True
	paragraph_shop.paragraph_format.space_after = 0
	paragraph_address = doc.add_paragraph(address)
	paragraph_address.paragraph_format.space_after = 0
	paragraph_mobile = doc.add_paragraph(f'{mobile}')
	subject = doc.add_paragraph("Sub: Bill for Supply of Aluminium light box.")

	#----Creating the table-----
	#Table Top Heading------------------------------------
	table = doc.add_table(rows=0, cols=0)
	table.style = "Table Grid"
	table.add_column(width=Pt(45))
	table.add_column(width=Pt(275))
	table.add_column(width=Pt(85))
	table.add_column(width=Pt(85))
	table.add_column(width=Pt(85))
	table.add_column(width=Pt(85))
	table.add_column(width=Pt(85))
	table.add_column(width=Pt(115))
	row1 = table.add_row().cells
	#_____________
	row1[0].text = "Sl No"
	row1[1].text = "Particle"
	row1[2].text = "Size"
	row1[3].text = "Sft"
	row1[4].text = "Quantity"
	row1[5].text = "Total Sft"
	row1[6].text = "rate"
	row1[7].text = "Amount"

	p0 = row1[0].paragraphs[0]
	p1 = row1[1].paragraphs[0]
	p2 = row1[2].paragraphs[0]
	p3 = row1[3].paragraphs[0]
	p4 = row1[4].paragraphs[0]
	p5 = row1[5].paragraphs[0]
	p6 = row1[6].paragraphs[0]
	p7 = row1[7].paragraphs[0]

	p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
	p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	p6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	p7.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	#----Getting the signboard specs------
	Amounts = []
	i = 1
	board_row = board.split('\n')
	for cells in board_row:
		cell_list = cells.split(" ")
		cell_int = [float(a) for a in cell_list]
		width0 = cell_int[0]
		height0 = cell_int[1]
		quantity = int(cell_int[2])
		rate = int(cell_int[3])
		if width0 % 1 == 0:
			width = int(width0)
		else:
			width = width0
			
		if height0 % 1 == 0:
			height = int(height0)
		else:
			height = height0

		#Particle
		if rate == 95:
			particle = "Pana Change & Servicing"
		else:
			particle = "New Sign Board"

		#Size
		size = str(width) + "x" + str(height)

		#Squarfoots
		sft0 = width * height

		if sft0 % 1 == 0:
			sft0 = int(sft0)

		sft = str(sft0)

		#Quantity
		Qty = "0" + str(quantity)

		#Total Squarefoot
		total_sft0 = sft0 * quantity
		total_sft = str(total_sft0)

		#Rate
		rat = str(rate)
		Rate = rat + "/-"

		#Amount
		amount0 = total_sft0 * rate
		if amount0 % 1 >= 0.5:
			amount_final = int(amount0) + 1
		else:
			amount_final = int(amount0) 
		amount = str(amount_final) + "/-"
		Amounts.append(amount_final)
		#Table--------
		row2 = table.add_row().cells

		row2[0].text = str(i)
		row2[1].text = particle
		row2[2].text = size
		row2[3].text = sft
		row2[4].text = Qty
		row2[5].text = total_sft
		row2[6].text = Rate
		row2[7].text = amount

		q0 = row2[0].paragraphs[0]
		q1 = row2[1].paragraphs[0]
		q2 = row2[2].paragraphs[0]
		q3 = row2[3].paragraphs[0]
		q4 = row2[4].paragraphs[0]
		q5 = row2[5].paragraphs[0]
		q6 = row2[6].paragraphs[0]
		q7 = row2[7].paragraphs[0]

		q0.alignment = WD_ALIGN_PARAGRAPH.LEFT
		q1.alignment = WD_ALIGN_PARAGRAPH.LEFT
		q2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		q3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		q4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		q5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		q6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		q7.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		i += 1

	#Creating Table lower section Structure-------------------
	row3 = table.add_row().cells
	row4 = table.add_row().cells
	row5 = table.add_row().cells
	row6 = table.add_row().cells
	row7 = table.add_row().cells

	a1 = table.cell(i,1)
	a2 = table.cell(i,6)
	b1 = table.cell(i + 1,1)
	b2 = table.cell(i + 1,6)
	c1 = table.cell(i + 2,1)
	c2 = table.cell(i + 2,6)
	d1 = table.cell(i + 3,1)
	d2 = table.cell(i + 3,6)
	e1 = table.cell(i + 4,1)
	e2 = table.cell(i + 4,6)
	A = a1.merge(a2)
	B = b1.merge(b2)
	C = c1.merge(c2)
	D = d1.merge(d2)
	E = e1.merge(e2)

	row3[1].text = "Total Amount ="
	row4[1].text = "(+)Vat@15% ="
	row5[1].text = " "
	row6[1].text = "(-)TDS@5% ="
	row7[1].text = "Grand Total ="

	r0 = row3[1].paragraphs[0]
	r1 = row4[1].paragraphs[0]
	r2 = row5[1].paragraphs[0]
	r3 = row6[1].paragraphs[0]
	r4 = row7[1].paragraphs[0]

	r0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	r1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	r2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	r3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	r4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	#Calculating The Lower Table Values

	AMOUNT = sum(Amounts)
	AMOUNT_wholeNumber = AMOUNT // 1
	AMOUNT_fractionPart = AMOUNT % 1

	if AMOUNT_fractionPart >= 0.5:
		totalAmount = AMOUNT_wholeNumber + 1
	else:
		totalAmount = AMOUNT_wholeNumber

	totalAmount_int = int(totalAmount)
	totalAmount_str = str(totalAmount_int) + "/-"
		


	VAT = totalAmount_int * 0.15

	VAT_wholeNumber = VAT // 1
	VAT_fractionPart = VAT % 1

	if VAT_fractionPart >= 0.5:
		vat = VAT_wholeNumber + 1
	else:
		vat = VAT_wholeNumber

	vat_int = int(vat)
	vat_str = str(vat_int) + "/-"


	IncludingVat_int = totalAmount_int + vat_int
	IncludingVat_str = str(IncludingVat_int) + "/-"


	tds = totalAmount_int * 0.05
	tds_wholeNumber = tds // 1
	tds_fractionPart = tds % 1

	if tds_fractionPart >= 0.5:
		TDS = tds_wholeNumber + 1
	else:
		TDS = tds_wholeNumber

	TDS_int = int(TDS)
	TDS_str = str(TDS_int) + "/-"
	



	GrandTotal = IncludingVat_int - TDS_int
	GrandTotal_str = str(GrandTotal) + "/-"
		


	#Inserting Table Lower section values
	row3[7].text = totalAmount_str
	row4[7].text = vat_str
	row5[7].text = IncludingVat_str
	row6[7].text = TDS_str
	row7[7].text = GrandTotal_str

	s0 = row3[7].paragraphs[0]
	s1 = row4[7].paragraphs[0]
	s2 = row5[7].paragraphs[0]
	s3 = row6[7].paragraphs[0]
	s4 = row7[7].paragraphs[0]

	s0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	s1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	s2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	s3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	s4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	#Bottom Section-----------------------------------------------
	doc.add_paragraph("\nIn Words: ")
	doc.add_paragraph("Thank you")
	doc.add_paragraph("\n")
	doc.add_paragraph("Grameen Media House")

	return doc
